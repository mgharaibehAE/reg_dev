import streamlit as st
import openai
import requests
from io import BytesIO
from docx import Document
import PyPDF2
import time
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
import google.generativeai as genai
import io
import tempfile


# Streamlit configuration
st.set_page_config(page_title="Cleco Regulatory Assistant", page_icon="🤖", layout="centered")

# Constants from secrets
OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
ASSISTANT_ID = st.secrets["ASSISTANT_ID"]
PASSWORD = st.secrets["login"]["password"]
GITHUB_API_URL = "https://api.github.com/repos/mgharaibehAE/assistant/contents/docs"
GITHUB_TOKEN = st.secrets["github"]["token"]
GROK_API_KEY = st.secrets["GROK_API_KEY"]
GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]

genai.configure(api_key=GEMINI_API_KEY)

# Sidebar
with st.sidebar:
    st.markdown("""
    **Disclaimer:** Regulatory Assistant can make mistakes. Verify important information.
    """)
    st.markdown("""
    **Instructions:**
    - Clearly enter your queries.
    - Use "Clear Chat" to reset the conversation.
    - Export chat history if required.
    """)
    if st.button("Clear Chat"):
        for key in ["messages", "thread_id", "authenticated", "file_thread_id", "file_chat_messages"]:
            st.session_state.pop(key, None)
        st.rerun()

# Authentication
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    pwd_input = st.text_input("Enter Password", type="password")
    if st.button("Login"):
        if pwd_input == PASSWORD:
            st.session_state.authenticated = True
            st.success("Login successful! Refreshing...")
            time.sleep(1)
            st.rerun()
        else:
            st.error("Incorrect password")
    st.stop()

st.title("Cleco Regulatory Assistant")

# OpenAI setup
openai.api_key = OPENAI_API_KEY

# Tabs for chat, document summary, and document chat
tab_chat, tab_docs, tab_upload = st.tabs(["Chat Assistant", "Document Summary", "Chat with Document"])

# First Tab (General Chat)
with tab_chat:
    if "messages" not in st.session_state:
        st.session_state.messages = []

    if "thread_id" not in st.session_state:
        thread = openai.beta.threads.create()
        st.session_state.thread_id = thread.id

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("Ask your question..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        openai.beta.threads.messages.create(
            thread_id=st.session_state.thread_id,
            role="user",
            content=prompt
        )

        run = openai.beta.threads.runs.create(
            thread_id=st.session_state.thread_id,
            assistant_id=ASSISTANT_ID
        )

        with st.spinner("Assistant is typing..."):
            while run.status not in ("completed", "failed"):
                time.sleep(1)
                run = openai.beta.threads.runs.retrieve(
                    thread_id=st.session_state.thread_id,
                    run_id=run.id
                )

        if run.status == "completed":
            messages = openai.beta.threads.messages.list(thread_id=st.session_state.thread_id)
            response = next((msg.content[0].text.value for msg in messages.data if msg.role == "assistant"), "")

            st.session_state.messages.append({"role": "assistant", "content": response})

            with st.chat_message("assistant"):
                st.markdown(response)
        else:
            st.error("Assistant failed to respond. Please retry.")

    chat_history = "\n".join(f"{msg['role'].capitalize()}: {msg['content']}" for msg in st.session_state.messages)
    st.download_button("Export Chat History", chat_history, file_name="chat_history.txt")

# Document Summary Tab
with tab_docs:
    headers = {"Authorization": f"token {GITHUB_TOKEN}"}
    response = requests.get(GITHUB_API_URL, headers=headers)

    if response.status_code == 200:
        files = response.json()
        doc_files = [file['name'] for file in files if file['name'].endswith(".docx")]

        selected_file = st.selectbox("Choose a Document", doc_files)
        if st.button("Go to Summary"):
            file_url = next((file['download_url'] for file in files if file['name'] == selected_file), None)
            if file_url:
                file_content = requests.get(file_url).content
                document = Document(BytesIO(file_content))
                doc_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
                st.markdown(f"### Content of {selected_file}")
                st.write(doc_text)
            else:
                st.error("Document not found.")
    else:
        st.error(f"Failed to load documents. Error: {response.status_code}, {response.json().get('message')}")


with tab_upload:
    st.header("Upload and Chat with Document")

    ai_model = st.selectbox("Choose AI Model", ["OpenAI Assistant", "Grok", "Gemini"])

    uploaded_file = st.file_uploader("Upload a Word (.docx) or PDF (.pdf) file", type=["docx", "pdf"], accept_multiple_files=True)

    if uploaded_file:
        uploaded_file_names = [file.name for file in uploaded_file]
        if 'last_uploaded_file_name' not in st.session_state or st.session_state.last_uploaded_file_name != uploaded_file_names:
            st.session_state.file_chat_messages = []
            st.session_state.file_thread_id = None
            st.session_state.last_uploaded_file_name = uploaded_file_names

        file_text = ""
        if ai_model != "Gemini":
            for file in uploaded_file:
                file_bytes = file.read()
                if file.type == "application/pdf":
                    images = convert_from_bytes(file_bytes)
                    for image in images:
                        text = pytesseract.image_to_string(image)
                        file_text += text + "\n"
                else:
                    doc = Document(file)
                    file_text += "\n".join(paragraph.text for paragraph in doc.paragraphs) + "\n"

        if "file_thread_id" not in st.session_state or st.session_state.file_thread_id is None:
            if ai_model == "OpenAI Assistant":
                openai.api_key = OPENAI_API_KEY
                file_thread = openai.beta.threads.create()
                st.session_state.file_thread_id = file_thread.id
                openai.beta.threads.messages.create(
                    thread_id=st.session_state.file_thread_id,
                    role="user",
                    content=f"The following document content is provided for context:\n\n{file_text}"
                )
            else:
                st.session_state.file_chat_context = file_text

        for message in st.session_state.file_chat_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        user_input = st.chat_input("Ask questions about the uploaded document(s)...")

        if user_input:
            st.session_state.file_chat_messages.append({"role": "user", "content": user_input})
            response = ""

            with st.spinner("Assistant is typing..."):
                if ai_model == "OpenAI Assistant":
                    openai.api_key = OPENAI_API_KEY
                    openai.beta.threads.messages.create(
                        thread_id=st.session_state.file_thread_id,
                        role="user",
                        content=user_input
                    )
                    run = openai.beta.threads.runs.create(
                        thread_id=st.session_state.file_thread_id,
                        assistant_id=ASSISTANT_ID
                    )
                    while run.status not in ("completed", "failed"):
                        time.sleep(1)
                        run = openai.beta.threads.runs.retrieve(
                            thread_id=st.session_state.file_thread_id,
                            run_id=run.id
                        )
                    if run.status == "completed":
                        messages = openai.beta.threads.messages.list(thread_id=st.session_state.file_thread_id)
                        response = next((msg.content[0].text.value for msg in messages.data if msg.role == "assistant"), "")
                    else:
                        response = "Assistant failed to respond. Please retry."

                elif ai_model == "Grok":
                    headers = {
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {GROK_API_KEY}"
                    }
                    data = {
                        "messages": [
                            {"role": "system", "content": "You are an assistant helping with document queries."},
                            {"role": "user", "content": st.session_state.file_chat_context + "\n" + user_input}
                        ],
                        "model": "grok-3-latest",
                        "stream": False,
                        "temperature": 0
                    }
                    grok_response = requests.post("https://api.x.ai/v1/chat/completions", headers=headers, json=data)
                    response = grok_response.json()['choices'][0]['message']['content']

                elif ai_model == "Gemini":
                    uploaded_files = []
                    for file in uploaded_file:
                        file_bytes = file.read()
                        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file.name.split('.')[-1]}") as tmp:
                            tmp.write(file_bytes)
                            tmp_path = tmp.name

                        uploaded_file_obj = genai.upload_file(path=tmp_path, mime_type=file.type)
                        uploaded_files.append(uploaded_file_obj)
                        os.unlink(tmp_path)

                    model = genai.GenerativeModel('gemini-1.5-flash')
                    chat = model.start_chat(history=[
                        {"role": "user", "parts": uploaded_files},
                    ])
                    gemini_response = chat.send_message(user_input)
                    response = gemini_response.text

            st.session_state.file_chat_messages.append({"role": "assistant", "content": response})

            with st.chat_message("assistant"):
                st.markdown(response)
